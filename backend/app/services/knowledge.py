"""Knowledge Hub service — parsing, chunking, embeddings, retrieval, RAG.

Self-contained module covering the whole ingest → retrieve → RAG pipeline.
Heavy work happens on the background ingest worker (see ``ingest_worker.py``)
so the HTTP request returns 202 immediately.
"""
from __future__ import annotations

import hashlib
import io
import json
import re
import time
from datetime import datetime, timezone
from decimal import Decimal
from pathlib import Path
from typing import Iterator

import httpx
import numpy as np
from fastapi import UploadFile
from loguru import logger
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.core.config import settings
from app.core.exceptions import AppError, NotFoundError, ValidationFailed
from app.models.knowledge import (
    KnowledgeBase, KnowledgeChunk, KnowledgeDocument, KnowledgeQuery,
)
from app.schemas.knowledge import (
    OllamaModelOut, OllamaModelsOut, RagChatIn, RetrievedChunkOut, RetrieveOut,
)


# ---------------------------------------------------------------------------
# Storage paths
# ---------------------------------------------------------------------------
def _backend_root() -> Path:
    return Path(__file__).resolve().parents[2]


def _kb_storage_root() -> Path:
    root = _backend_root() / settings.knowledge_storage_dir
    root.mkdir(parents=True, exist_ok=True)
    return root


def _doc_dir(kb_id: str, doc_id: str) -> Path:
    d = _kb_storage_root() / kb_id / doc_id
    d.mkdir(parents=True, exist_ok=True)
    return d


# ---------------------------------------------------------------------------
# MIME / extension helpers
# ---------------------------------------------------------------------------
SUPPORTED_EXTS = {".pdf", ".docx", ".txt", ".md", ".markdown", ".html", ".htm"}


def _ext_from_name(name: str) -> str:
    return Path(name).suffix.lower()


def _is_supported(name: str) -> bool:
    return _ext_from_name(name) in SUPPORTED_EXTS


def _sha256(blob: bytes) -> str:
    return hashlib.sha256(blob).hexdigest()


# ---------------------------------------------------------------------------
# Parsers — return (full_text, blocks) where blocks is a list of
# (page_number_or_None, char_start, char_end, text) for chunk metadata.
# ---------------------------------------------------------------------------
ParsedBlock = tuple[int | None, int, int, str]


def _parse_pdf(path: Path) -> tuple[str, list[ParsedBlock]]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    full_parts: list[str] = []
    blocks: list[ParsedBlock] = []
    cursor = 0
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = _normalize_whitespace(text)
        if not text.strip():
            continue
        start = cursor
        full_parts.append(text)
        cursor += len(text) + 2  # +2 for the "\n\n" joiner
        blocks.append((i, start, start + len(text), text))
    full_text = "\n\n".join(full_parts)
    return full_text, blocks


def _parse_docx(path: Path) -> tuple[str, list[ParsedBlock]]:
    from docx import Document

    doc = Document(str(path))
    full_parts: list[str] = []
    blocks: list[ParsedBlock] = []
    cursor = 0
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        start = cursor
        full_parts.append(text)
        cursor += len(text) + 2
        blocks.append((None, start, start + len(text), text))
    # Also pull table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = (cell.text or "").strip()
                if not text:
                    continue
                start = cursor
                full_parts.append(text)
                cursor += len(text) + 2
                blocks.append((None, start, start + len(text), text))
    full_text = "\n\n".join(full_parts)
    return full_text, blocks


def _parse_html(path: Path) -> tuple[str, list[ParsedBlock]]:
    from bs4 import BeautifulSoup

    raw = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(raw, "html.parser")
    for s in soup(["script", "style", "noscript"]):
        s.decompose()
    text = soup.get_text(separator="\n\n")
    text = _normalize_whitespace(text)
    return text, _paragraph_blocks(text)


def _parse_text(path: Path) -> tuple[str, list[ParsedBlock]]:
    raw = path.read_text(encoding="utf-8", errors="replace")
    text = _normalize_whitespace(raw)
    return text, _paragraph_blocks(text)


def _parse_markdown(path: Path) -> tuple[str, list[ParsedBlock]]:
    # Treat markdown source as text for retrieval — keeps source-style chunks
    # easier to cite back. (Rendered HTML chunks would lose the headings.)
    return _parse_text(path)


def _paragraph_blocks(text: str) -> list[ParsedBlock]:
    blocks: list[ParsedBlock] = []
    cursor = 0
    for para in re.split(r"\n\s*\n", text):
        para = para.strip()
        if not para:
            cursor += 2
            continue
        start = cursor
        blocks.append((None, start, start + len(para), para))
        cursor = start + len(para) + 2
    return blocks


def _normalize_whitespace(text: str) -> str:
    # Collapse runs of whitespace inside lines but preserve paragraph breaks.
    out_lines: list[str] = []
    for line in text.splitlines():
        cleaned = re.sub(r"[ \t]+", " ", line).strip()
        out_lines.append(cleaned)
    # Collapse 3+ blank lines into 2
    text = "\n".join(out_lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def parse_document(path: Path, mime: str | None, name: str) -> tuple[str, list[ParsedBlock]]:
    """Dispatch on extension first, mime as fallback."""
    ext = _ext_from_name(name) or _ext_from_name(path.name)
    if ext == ".pdf":
        return _parse_pdf(path)
    if ext == ".docx":
        return _parse_docx(path)
    if ext in {".html", ".htm"}:
        return _parse_html(path)
    if ext in {".md", ".markdown"}:
        return _parse_markdown(path)
    # Default to plain text
    return _parse_text(path)


# ---------------------------------------------------------------------------
# Chunker
# ---------------------------------------------------------------------------
def chunk_text(
    blocks: list[ParsedBlock], *, chunk_size: int, chunk_overlap: int
) -> list[dict]:
    """Sliding-window chunker that respects paragraph boundaries.

    Returns a list of dicts: ``{ordinal, page_number, char_start, char_end, text}``.
    """
    if chunk_size <= 0:
        raise ValueError("chunk_size must be positive")
    if chunk_overlap < 0 or chunk_overlap >= chunk_size:
        raise ValueError("chunk_overlap must satisfy 0 <= overlap < chunk_size")

    chunks: list[dict] = []
    buffer: list[ParsedBlock] = []
    buffer_len = 0

    def _emit():
        if not buffer:
            return
        text_parts = [b[3] for b in buffer]
        text = "\n\n".join(text_parts)
        if not text.strip():
            return
        page_number = buffer[0][0]  # take the first block's page
        char_start = buffer[0][1]
        char_end = buffer[-1][2]
        chunks.append({
            "ordinal": len(chunks),
            "page_number": page_number,
            "char_start": char_start,
            "char_end": char_end,
            "text": text,
        })

    for blk in blocks:
        _, _, _, btext = blk
        # If a single block is itself larger than chunk_size, slice it.
        if len(btext) > chunk_size and not buffer:
            _slice_oversized_block(blk, chunk_size, chunk_overlap, chunks)
            continue
        if buffer_len + len(btext) + 2 > chunk_size and buffer:
            _emit()
            # Build overlap: keep the tail of the last buffer up to chunk_overlap chars
            if chunk_overlap > 0:
                tail_text = "\n\n".join(b[3] for b in buffer)
                if len(tail_text) > chunk_overlap:
                    tail_text = tail_text[-chunk_overlap:]
                if tail_text.strip():
                    # Synthesize a single block representing the overlap window
                    # using the LAST block's coordinates so the citation still
                    # points at a valid range.
                    last = buffer[-1]
                    buffer = [(last[0], last[2] - len(tail_text), last[2], tail_text)]
                    buffer_len = len(tail_text)
                else:
                    buffer = []
                    buffer_len = 0
            else:
                buffer = []
                buffer_len = 0
        buffer.append(blk)
        buffer_len += len(btext) + 2

    _emit()
    return chunks


def _slice_oversized_block(
    block: ParsedBlock, chunk_size: int, chunk_overlap: int, out: list[dict]
) -> None:
    page, start, _end, text = block
    stride = chunk_size - chunk_overlap
    pos = 0
    while pos < len(text):
        slice_text = text[pos: pos + chunk_size]
        if not slice_text.strip():
            break
        out.append({
            "ordinal": len(out),
            "page_number": page,
            "char_start": start + pos,
            "char_end": start + pos + len(slice_text),
            "text": slice_text,
        })
        if pos + chunk_size >= len(text):
            break
        pos += stride


def count_tokens(text: str) -> int:
    """Approximate token count via tiktoken. Falls back to char/4 heuristic
    if the library isn't installed."""
    try:
        import tiktoken
        enc = tiktoken.get_encoding("cl100k_base")
        return len(enc.encode(text, disallowed_special=()))
    except Exception:
        return max(1, len(text) // 4)


# ---------------------------------------------------------------------------
# Embeddings
# ---------------------------------------------------------------------------
def _hash_pseudo_embedding(text: str, dim: int) -> np.ndarray:
    """Deterministic, terrible-quality fallback so the pipeline never hard-fails.

    Mixes a bag-of-character-trigrams hash into a fixed-dim vector. Loud
    ``WARN`` when used — production users should run Ollama with nomic-embed-text.
    """
    vec = np.zeros(dim, dtype=np.float32)
    s = text.lower()
    for i in range(len(s) - 2):
        tri = s[i:i + 3]
        h = int(hashlib.md5(tri.encode("utf-8")).hexdigest()[:8], 16)
        vec[h % dim] += 1.0
    norm = float(np.linalg.norm(vec))
    if norm > 0:
        vec /= norm
    return vec


def _embed_ollama(texts: list[str], model: str) -> list[np.ndarray]:
    out: list[np.ndarray] = []
    with httpx.Client(timeout=120.0) as client:
        for text in texts:
            try:
                r = client.post(
                    f"{settings.ollama_host}/api/embeddings",
                    json={"model": model, "prompt": text},
                )
                r.raise_for_status()
                vec = np.asarray(r.json().get("embedding") or [], dtype=np.float32)
                if vec.size == 0:
                    raise ValueError("empty embedding")
                n = float(np.linalg.norm(vec))
                if n > 0:
                    vec /= n
                out.append(vec)
            except Exception as e:
                logger.warning("Ollama embed failed for chunk ({}); using hash fallback", e)
                out.append(_hash_pseudo_embedding(text, dim=768))
    return out


def _embed_openai(texts: list[str], model: str, *, api_key_override: str | None = None) -> list[np.ndarray]:
    key = api_key_override or settings.openai_api_key
    if not key:
        raise AppError("OpenAI API key not configured", code="ai_provider_not_configured")
    from openai import OpenAI

    client = OpenAI(api_key=key)
    out: list[np.ndarray] = []
    # OpenAI supports batch up to ~2048 inputs; chunk to be safe.
    BATCH = 96
    for i in range(0, len(texts), BATCH):
        batch = texts[i: i + BATCH]
        resp = client.embeddings.create(model=model, input=batch)
        for row in resp.data:
            vec = np.asarray(row.embedding, dtype=np.float32)
            n = float(np.linalg.norm(vec))
            if n > 0:
                vec /= n
            out.append(vec)
    return out


def embed(texts: list[str], *, provider: str, model: str,
          api_key_override: str | None = None) -> list[np.ndarray]:
    if not texts:
        return []
    if provider == "ollama":
        try:
            return _embed_ollama(texts, model)
        except Exception as e:
            logger.warning("Ollama unreachable ({}); falling back to hash embeddings", e)
            return [_hash_pseudo_embedding(t, dim=768) for t in texts]
    if provider == "openai":
        return _embed_openai(texts, model, api_key_override=api_key_override)
    raise AppError(f"Unsupported embedding provider: {provider}", code="bad_provider")


# ---------------------------------------------------------------------------
# Vector store (numpy + SQLite BLOB)
# ---------------------------------------------------------------------------
def encode_vec(vec: np.ndarray) -> bytes:
    return np.asarray(vec, dtype=np.float32).tobytes()


def decode_vec(blob: bytes) -> np.ndarray:
    return np.frombuffer(blob, dtype=np.float32)


def cosine_topk(query: np.ndarray, matrix: np.ndarray, k: int) -> list[tuple[int, float]]:
    """Cosine similarity, assuming both query and matrix rows are
    L2-normalized. Returns sorted [(row_index, score), ...] of length min(k, n)."""
    if matrix.size == 0:
        return []
    scores = matrix @ query
    k = min(int(k), scores.shape[0])
    if k <= 0:
        return []
    idx = np.argpartition(-scores, k - 1)[:k]
    idx = idx[np.argsort(-scores[idx])]
    return [(int(i), float(scores[i])) for i in idx]


# ---------------------------------------------------------------------------
# Ingest pipeline
# ---------------------------------------------------------------------------
def _validate_upload_size(byte_size: int) -> None:
    limit = int(settings.knowledge_max_upload_mb) * 1024 * 1024
    if byte_size > limit:
        raise ValidationFailed(
            f"File too large ({byte_size / 1024 / 1024:.1f} MB > {settings.knowledge_max_upload_mb} MB limit)"
        )


async def ingest_upload(
    db: Session, *, kb_id: str, upload: UploadFile, user_id: str | None
) -> KnowledgeDocument:
    name = upload.filename or "upload"
    if not _is_supported(name):
        raise ValidationFailed(
            f"Unsupported file type: {_ext_from_name(name) or '(none)'}. "
            f"Supported: {sorted(SUPPORTED_EXTS)}"
        )
    data = await upload.read()
    _validate_upload_size(len(data))
    return _create_doc_from_blob(
        db, kb_id=kb_id, name=name, source_type="upload",
        source_ref=name, mime_type=upload.content_type, blob=data,
    )


def ingest_paste(
    db: Session, *, kb_id: str, name: str, text: str, user_id: str | None
) -> KnowledgeDocument:
    safe_name = (name or "Pasted note").strip()
    if not safe_name.lower().endswith(".txt"):
        safe_name = f"{safe_name}.txt"
    data = text.encode("utf-8")
    _validate_upload_size(len(data))
    return _create_doc_from_blob(
        db, kb_id=kb_id, name=safe_name, source_type="paste",
        source_ref=None, mime_type="text/plain", blob=data,
    )


def ingest_url(
    db: Session, *, kb_id: str, url: str, name: str | None, user_id: str | None
) -> KnowledgeDocument:
    timeout = int(settings.knowledge_url_fetch_timeout)
    try:
        with httpx.Client(timeout=timeout, follow_redirects=True) as client:
            resp = client.get(url, headers={"User-Agent": "EnterpriseCoreAI/0.1 (+knowledge-hub)"})
            resp.raise_for_status()
    except httpx.HTTPError as e:
        raise AppError(f"Could not fetch URL: {e}", code="url_fetch_failed", status_code=502) from e
    content_type = resp.headers.get("content-type", "").split(";")[0].strip() or "text/html"
    data = resp.content
    _validate_upload_size(len(data))
    fallback_name = (name or url.split("/")[-1] or "fetched-page").strip()
    # Guess the right extension from content-type if filename has none
    if not _ext_from_name(fallback_name):
        if "pdf" in content_type:
            fallback_name += ".pdf"
        elif "html" in content_type or "xml" in content_type:
            fallback_name += ".html"
        else:
            fallback_name += ".txt"
    return _create_doc_from_blob(
        db, kb_id=kb_id, name=fallback_name, source_type="url",
        source_ref=url, mime_type=content_type, blob=data,
    )


def _create_doc_from_blob(
    db: Session, *, kb_id: str, name: str, source_type: str,
    source_ref: str | None, mime_type: str | None, blob: bytes,
) -> KnowledgeDocument:
    kb = db.get(KnowledgeBase, kb_id)
    if not kb:
        raise NotFoundError("Knowledge base not found")
    doc = KnowledgeDocument(
        kb_id=kb_id, name=name, source_type=source_type, source_ref=source_ref,
        mime_type=mime_type, byte_size=len(blob), sha256=_sha256(blob),
        status="queued",
    )
    db.add(doc)
    db.flush()  # populate id
    target_dir = _doc_dir(kb_id, doc.id)
    ext = _ext_from_name(name) or ".bin"
    target = target_dir / f"original{ext}"
    target.write_bytes(blob)
    doc.storage_path = str(target)
    db.commit()
    db.refresh(doc)
    return doc


def delete_document_blob(doc: KnowledgeDocument) -> None:
    """Best-effort blob cleanup. Errors are logged but don't block delete."""
    try:
        if doc.storage_path:
            p = Path(doc.storage_path)
            if p.exists():
                p.unlink()
            # Also remove the extracted cache + parent dir if now empty
            parent = p.parent
            extracted = parent / "extracted.txt"
            if extracted.exists():
                extracted.unlink()
            try:
                parent.rmdir()
            except OSError:
                pass  # not empty
    except Exception as e:
        logger.warning("Failed to remove blob for doc {}: {}", doc.id, e)


def requeue_document(db: Session, doc: KnowledgeDocument, *, reset_chunks: bool = True) -> None:
    if reset_chunks:
        db.query(KnowledgeChunk).filter(KnowledgeChunk.document_id == doc.id).delete()
    doc.status = "queued"
    doc.error_message = None
    doc.chunk_count = 0
    doc.ingested_at = None
    db.commit()


def requeue_kb(db: Session, kb_id: str) -> int:
    docs = db.scalars(
        select(KnowledgeDocument).where(KnowledgeDocument.kb_id == kb_id)
    ).all()
    for doc in docs:
        requeue_document(db, doc, reset_chunks=True)
    return len(docs)


# ---------------------------------------------------------------------------
# Worker — process one queued doc end-to-end
# ---------------------------------------------------------------------------
def _set_status(db: Session, doc: KnowledgeDocument, status: str,
                error: str | None = None) -> None:
    doc.status = status
    if error is not None:
        doc.error_message = error
    db.commit()


def process_document(db: Session, doc_id: str) -> bool:
    """Parse → chunk → embed → ready. Returns True on success."""
    doc = db.get(KnowledgeDocument, doc_id)
    if not doc:
        return False
    if doc.status not in {"queued", "failed"}:
        return False
    kb = db.get(KnowledgeBase, doc.kb_id)
    if not kb:
        _set_status(db, doc, "failed", "Knowledge base not found")
        return False

    try:
        _set_status(db, doc, "parsing")
        path = Path(doc.storage_path or "")
        if not path.exists():
            raise FileNotFoundError(f"Original blob missing: {path}")

        full_text, blocks = parse_document(path, doc.mime_type, doc.name)
        if not full_text.strip():
            raise ValueError("Document has no extractable text")

        extracted_cache = path.parent / "extracted.txt"
        try:
            extracted_cache.write_text(full_text, encoding="utf-8")
        except Exception as e:
            logger.warning("Could not write extracted cache for {}: {}", doc.id, e)

        doc.char_count = len(full_text)
        doc.page_count = len({b[0] for b in blocks if b[0] is not None})

        chunk_dicts = chunk_text(
            blocks, chunk_size=kb.chunk_size, chunk_overlap=kb.chunk_overlap
        )
        if not chunk_dicts:
            raise ValueError("Chunker produced 0 chunks")

        _set_status(db, doc, "embedding")
        texts = [c["text"] for c in chunk_dicts]
        vectors = embed(
            texts, provider=kb.embedding_provider, model=kb.embedding_model
        )
        if len(vectors) != len(chunk_dicts):
            raise ValueError(
                f"Embedding count mismatch: {len(vectors)} vs {len(chunk_dicts)}"
            )

        # Bulk insert chunks
        for c, vec in zip(chunk_dicts, vectors):
            db.add(KnowledgeChunk(
                document_id=doc.id, kb_id=doc.kb_id,
                ordinal=c["ordinal"], text=c["text"],
                page_number=c["page_number"],
                char_start=c["char_start"], char_end=c["char_end"],
                token_count=count_tokens(c["text"]),
                embedding=encode_vec(vec),
                embedding_model=f"{kb.embedding_provider}/{kb.embedding_model}",
            ))
        doc.chunk_count = len(chunk_dicts)
        doc.ingested_at = datetime.now(timezone.utc)
        doc.status = "ready"
        doc.error_message = None
        db.commit()
        try:
            from app.services.event_bus import publish_event

            publish_event(
                "knowledge.document.ingested",
                payload={
                    "document_id": doc.id,
                    "knowledge_base_id": doc.knowledge_base_id,
                    "chunk_count": doc.chunk_count,
                },
                tenant_id=doc.tenant_id,
            )
        except Exception:  # pragma: no cover
            pass
        return True
    except Exception as e:
        logger.exception("Ingest failed for doc {}: {}", doc_id, e)
        try:
            db.rollback()
            doc = db.get(KnowledgeDocument, doc_id)
            if doc:
                doc.status = "failed"
                doc.error_message = str(e)[:1000]
                db.commit()
        except Exception:
            pass
        return False


def process_pending(db: Session, *, limit: int = 5) -> int:
    """Process up to ``limit`` queued documents in this pass."""
    docs = db.scalars(
        select(KnowledgeDocument)
        .where(KnowledgeDocument.status == "queued")
        .order_by(KnowledgeDocument.created_at)
        .limit(limit)
    ).all()
    done = 0
    for doc in docs:
        if process_document(db, doc.id):
            done += 1
    return done


# ---------------------------------------------------------------------------
# Retrieval
# ---------------------------------------------------------------------------
def _load_kb_vectors(db: Session, kb_ids: list[str]) -> tuple[list[KnowledgeChunk], np.ndarray]:
    chunks = db.scalars(
        select(KnowledgeChunk).where(KnowledgeChunk.kb_id.in_(kb_ids))
    ).all()
    chunks = [c for c in chunks if c.embedding]
    if not chunks:
        return [], np.zeros((0, 0), dtype=np.float32)
    matrix = np.stack([decode_vec(c.embedding) for c in chunks])
    return chunks, matrix


def retrieve(
    db: Session, *, kb_ids: list[str], query: str, top_k: int = 6,
    user_id: str | None = None,
) -> RetrieveOut:
    if not kb_ids:
        raise ValidationFailed("At least one KB must be selected")
    kbs = db.scalars(select(KnowledgeBase).where(KnowledgeBase.id.in_(kb_ids))).all()
    if not kbs:
        raise NotFoundError("No knowledge bases found")
    kb_name_by_id = {kb.id: kb.name for kb in kbs}
    primary = kbs[0]
    started = time.time()
    q_vecs = embed([query], provider=primary.embedding_provider, model=primary.embedding_model)
    if not q_vecs:
        return RetrieveOut(
            query=query, chunks=[],
            embedding_provider=primary.embedding_provider,
            embedding_model=primary.embedding_model,
            latency_ms=int((time.time() - started) * 1000),
        )
    qvec = q_vecs[0]
    chunks, matrix = _load_kb_vectors(db, kb_ids)
    if matrix.size == 0:
        return RetrieveOut(
            query=query, chunks=[],
            embedding_provider=primary.embedding_provider,
            embedding_model=primary.embedding_model,
            latency_ms=int((time.time() - started) * 1000),
        )

    # Vectors may differ in length if a KB mixes models (shouldn't happen, but
    # be defensive): keep only those matching qvec dim.
    if qvec.shape[0] != matrix.shape[1]:
        keep = [i for i, c in enumerate(chunks) if decode_vec(c.embedding).shape[0] == qvec.shape[0]]
        if not keep:
            return RetrieveOut(
                query=query, chunks=[],
                embedding_provider=primary.embedding_provider,
                embedding_model=primary.embedding_model,
                latency_ms=int((time.time() - started) * 1000),
            )
        chunks = [chunks[i] for i in keep]
        matrix = matrix[keep]

    ranked = cosine_topk(qvec, matrix, top_k * 2)  # over-fetch then de-dup
    seen_pages: set[tuple[str, int | None]] = set()
    out: list[RetrievedChunkOut] = []
    for idx, score in ranked:
        c = chunks[idx]
        key = (c.document_id, c.page_number)
        if key in seen_pages:
            continue
        seen_pages.add(key)
        doc = db.get(KnowledgeDocument, c.document_id)
        if not doc:
            continue
        out.append(RetrievedChunkOut(
            chunk_id=c.id, document_id=c.document_id,
            document_name=doc.name, kb_id=c.kb_id,
            kb_name=kb_name_by_id.get(c.kb_id, ""),
            text=c.text, page_number=c.page_number, score=score,
        ))
        if len(out) >= top_k:
            break

    latency_ms = int((time.time() - started) * 1000)
    return RetrieveOut(
        query=query, chunks=out,
        embedding_provider=primary.embedding_provider,
        embedding_model=primary.embedding_model,
        latency_ms=latency_ms,
    )


# ---------------------------------------------------------------------------
# RAG chat (streaming SSE)
# ---------------------------------------------------------------------------
def build_rag_prompt(question: str, chunks: list[RetrievedChunkOut]) -> tuple[str, str]:
    system = (
        "You are a precise assistant for the EnterpriseCore AI Suite. "
        "Answer the question using ONLY the sources below. "
        "Cite each claim with bracketed numbers like [1] [2] that match the "
        "source index. If the sources don't contain enough information, say so "
        "plainly — don't invent details."
    )
    src_lines = []
    for i, c in enumerate(chunks, start=1):
        head = f"[{i}] (KB: {c.kb_name} / Doc: {c.document_name}"
        if c.page_number is not None:
            head += f" / Page: {c.page_number}"
        head += ")"
        src_lines.append(f"{head}\n{c.text}")
    sources_block = "\n\n".join(src_lines) if src_lines else "(no sources retrieved)"
    user = (
        "SOURCES:\n"
        f"{sources_block}\n\n"
        "QUESTION:\n"
        f"{question}"
    )
    return system, user


def _sse(event: str, payload: dict) -> bytes:
    return f"event: {event}\ndata: {json.dumps(payload, default=str)}\n\n".encode("utf-8")


def rag_chat_stream(*, db: Session, user_id: str | None, payload: RagChatIn) -> Iterator[bytes]:
    """Generator that yields SSE events for a streaming RAG response.

    Imports the streaming AI helper lazily to dodge circular import at startup.
    """
    from app.services import ai as ai_svc
    from app.models.ai import AiConversation, AiMessage

    question = payload.messages[-1].content if payload.messages else ""
    if not question.strip():
        yield _sse("error", {"detail": "Empty question"})
        yield _sse("done", {})
        return

    started = time.time()
    try:
        retrieve_out = retrieve(
            db, kb_ids=payload.kb_ids, query=question,
            top_k=payload.top_k, user_id=user_id,
        )
    except AppError as e:
        yield _sse("error", {"detail": e.message})
        yield _sse("done", {})
        return

    chunks = retrieve_out.chunks
    sources_payload = [{
        "index": i + 1,
        "chunk_id": c.chunk_id, "document_id": c.document_id,
        "document_name": c.document_name, "kb_id": c.kb_id, "kb_name": c.kb_name,
        "page_number": c.page_number, "score": c.score,
        "text": c.text[:500] + ("…" if len(c.text) > 500 else ""),
    } for i, c in enumerate(chunks)]
    yield _sse("sources", {"chunks": sources_payload})

    system, user = build_rag_prompt(question, chunks)
    # Persist conversation + user message (mirror /ai/chat behaviour)
    if payload.conversation_id:
        conv = db.get(AiConversation, payload.conversation_id)
        if not conv:
            yield _sse("error", {"detail": "Conversation not found"})
            yield _sse("done", {})
            return
    else:
        title = question[:80]
        conv = AiConversation(
            user_id=user_id, title=title,
            provider=payload.provider or "ollama",
            model=payload.model, module="knowledge",
        )
        db.add(conv)
        db.flush()
    for m in payload.messages:
        db.add(AiMessage(
            conversation_id=conv.id, role=m.role, content=m.content,
            tokens_in=0, tokens_out=0,
        ))
    db.commit()

    full_answer_parts: list[str] = []
    final_meta: dict = {}
    try:
        gen = ai_svc.stream_call(
            messages=[
                ai_svc.AiMessage(role="system", content=system),
                ai_svc.AiMessage(role="user", content=user),
            ],
            provider=payload.provider, model=payload.model,
            max_tokens=payload.max_tokens, temperature=payload.temperature,
            feature="rag_chat", db=db, user_id=user_id,
        )
        for ev_type, ev_payload in gen:
            if ev_type == "token":
                full_answer_parts.append(ev_payload.get("text", ""))
                yield _sse("token", ev_payload)
            elif ev_type == "usage":
                final_meta = ev_payload
                yield _sse("usage", {
                    **ev_payload,
                    "conversation_id": conv.id,
                    "retrieved_chunk_ids": [c.chunk_id for c in chunks],
                })
            elif ev_type == "error":
                yield _sse("error", ev_payload)
    except AppError as e:
        yield _sse("error", {"detail": e.message})

    answer_text = "".join(full_answer_parts)
    # Persist assistant message + audit row
    try:
        db.add(AiMessage(
            conversation_id=conv.id, role="assistant", content=answer_text,
            tokens_in=int(final_meta.get("tokens_in", 0)),
            tokens_out=int(final_meta.get("tokens_out", 0)),
        ))
        conv.updated_at = datetime.now(timezone.utc)
        db.add(KnowledgeQuery(
            user_id=user_id,
            kb_ids_json=json.dumps(payload.kb_ids),
            conversation_id=conv.id,
            question=question,
            retrieved_chunk_ids_json=json.dumps([c.chunk_id for c in chunks]),
            answer=answer_text,
            provider=final_meta.get("provider"),
            model=final_meta.get("model"),
            latency_ms=int((time.time() - started) * 1000),
            tokens_in=int(final_meta.get("tokens_in", 0)),
            tokens_out=int(final_meta.get("tokens_out", 0)),
            cost_usd=Decimal(str(final_meta.get("cost_usd", "0"))),
        ))
        db.commit()
    except Exception as e:
        logger.warning("Failed to persist RAG audit row: {}", e)

    yield _sse("done", {"conversation_id": conv.id})


# ---------------------------------------------------------------------------
# Ollama model manager
# ---------------------------------------------------------------------------
def list_ollama_models() -> OllamaModelsOut:
    host = settings.ollama_host
    try:
        with httpx.Client(timeout=5.0) as client:
            r = client.get(f"{host}/api/tags")
            r.raise_for_status()
            data = r.json()
    except Exception as e:
        logger.info("Ollama unreachable at {}: {}", host, e)
        return OllamaModelsOut(models=[], host=host, reachable=False)
    models_raw = data.get("models", []) or []
    models = []
    for m in models_raw:
        details = m.get("details") or {}
        models.append(OllamaModelOut(
            name=m.get("name") or m.get("model") or "?",
            size_bytes=int(m.get("size") or 0),
            modified_at=m.get("modified_at"),
            parameter_size=details.get("parameter_size"),
            family=details.get("family") or (details.get("families") or [None])[0],
        ))
    return OllamaModelsOut(models=models, host=host, reachable=True)


def pull_ollama_model_stream(model_name: str) -> Iterator[bytes]:
    host = settings.ollama_host
    if not model_name.strip():
        yield _sse("error", {"detail": "model name required"})
        yield _sse("done", {})
        return
    try:
        with httpx.Client(timeout=None) as client:
            with client.stream(
                "POST", f"{host}/api/pull",
                json={"name": model_name.strip(), "stream": True},
            ) as r:
                if r.status_code >= 400:
                    yield _sse("error", {"detail": f"Ollama returned HTTP {r.status_code}"})
                    yield _sse("done", {})
                    return
                for line in r.iter_lines():
                    if not line:
                        continue
                    try:
                        ev = json.loads(line)
                    except json.JSONDecodeError:
                        continue
                    yield _sse("progress", ev)
                    if ev.get("status") == "success":
                        break
    except Exception as e:
        yield _sse("error", {"detail": f"Pull failed: {e}"})
    yield _sse("done", {"model": model_name})
